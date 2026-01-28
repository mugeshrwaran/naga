import json
import streamlit as st
import google.generativeai as genai
import dotenv
import os
from docx import Document
from io import BytesIO
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter
from jsontostring import convert_sales_report_to_string


def parse_explicit_counts(data_series):
    """
    Parses a pandas Series of strings that may contain comma-separated items
    with explicit counts like 'Item - 100' or 'Item – 50'.
    Returns a dictionary of aggregated {item_name: count}.
    """
    total_counts = {}
    for value in data_series:
        if pd.isna(value) or str(value).strip().lower() == 'nan':
            continue
        
        # Split by comma for multiple items in one cell
        items_raw = [i.strip() for i in str(value).split(",") if i.strip()]
        
        for raw_item in items_raw:
            # Check for delimiters like '–' (long dash) or '-' (standard dash)
            if "–" in raw_item:
                parts = raw_item.split("–")
            elif "-" in raw_item:
                parts = raw_item.split("-")
            else:
                parts = [raw_item]
            
            name = parts[0].strip()
            
            # Extract count if present, otherwise default to 1
            if len(parts) > 1:
                try:
                    # Remove any non-numeric chars before converting
                    count_str = "".join(filter(str.isdigit, parts[1]))
                    count = int(count_str) if count_str else 1
                except ValueError:
                    count = 1
            else:
                count = 1
            
            total_counts[name] = total_counts.get(name, 0) + count
    return total_counts


# Load environment variables
dotenv.load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

st.logo(
    "Naga E-Store.png",
    size="large",
)

st.markdown(
"""
<style>
    header[data-testid="stHeader"]::after {
        content: "Naga Foods";
        font-size: 2.5rem;
        font-weight: 650;
        color: var(--text-color);
        position: absolute;
        left: 50%;
        transform: translateX(-50%);
    }
</style>
"""
, unsafe_allow_html=True)

def analyze_audio_with_gemini(audio_file):
    # Configure generation parameters for consistency and JSON mode
    generation_config = {
        "response_mime_type": "application/json"
    }
    
    model = genai.GenerativeModel(
        "gemini-2.5-pro",
        generation_config=generation_config
    )
    
    # Combined prompt for direct audio analysis
    analysis_prompt = """
CONFIGURATION

Manufacturer/Company: Naga
Sales Representative: Naga Foods salesperson

------------------------------------------------------------

CRITICAL INSTRUCTION - Brand Identification

1. Naga's OWN Products
- "Naga brand"
- "Our product" / "Our company's product"
- "Naga Foods product"
- If no brand is mentioned, assume it's Naga's
- Do not repeat products if already mentioned

2. Competitor Brands – ALL other brand names mentioned, including:
- Nandi, Sankar, Shakti, Aachi, MTR, Britannia, etc.

3. Online Retailers – Any online platforms mentioned (e.g., Amazon, Flipkart, BigBasket, etc.)

IMPORTANT: DO NOT assume a product is Naga's unless explicitly stated!

------------------------------------------------------------

SPEAKER CONTEXT RULES (CRITICAL)
 
Before mapping brands and products, determine who is speaking:
 
- If the **Sales Representative** mentions a product, assume it is **Naga's product** unless they clearly say it's a competitor.
- If the **Customer (store owner)** mentions a product or brand name, assume it is a **competitor brand**, unless the Sales Rep later confirms it belongs to Naga.
- If both speakers mention the same product name, assign ownership based on context and tone:
  - If Sales Rep is promoting or explaining → Naga's product.
  - If Customer is comparing or complaining → Competitor product.
- When uncertain, label it as **"Ambiguous – Needs context"** and do not count it in Naga product analysis.
------------------------------------------------------------

Brand & Product Mapping (Complete this FIRST)

Before analysis, categorize ALL brands and products mentioned in the conversation:

A. Naga Brand Products
- [List of Products from Manufacturer, Schemes are not discussed]
- [List of Products from Manufacturer, where schemes are discussed]

B. Competitor Brands Mentioned
- [List EACH competitor brand separately with details]

------------------------------------------------------------

Comprehensive Sales Analysis

Listen to this tamil audio conversation and provide analysis without transcribing first.

IMPORTANT: Start directly with the analysis content. Do not include introductory phrases.

------------------------------------------------------------

1. Conversation Summary
    - give a brief summary of the conversation (3-5 sentences)

------------------------------------------------------------

2. Sales Matrix

Naga Products Performance
- Naga products promoted: Which Naga products were pitched? Customer response?
- Volume pushed / upselling: Bulk orders or larger pack sizes attempted? Quantities?
- Schemes offered: Naga schemes, discounts, free-piece offers mentioned?
    [Give details of all the schemes mentioned with specifics like which product is offered with which scheme, discount %, which item is free for which scheme, etc.]
- Cross-selling within Naga portfolio: Were multiple Naga products bundled?
    [Give details of any cross-selling efforts]
- Acceptance/Rejection: Which Naga products did customer accept or reject?

Sales Barriers
- Objections raised: What prevented Naga sales?
- Competitor advantages cited: What specific advantages did competitors have?

------------------------------------------------------------
3. Customer Buying Patterns

A. Regularly buying products (Customer commits to buy BEFORE schemes are explained OR shows clear intent to buy regardless of schemes)
    - [List products where customer showed immediate interest or agreed to buy before any schemes/offers were mentioned]
    - [Also include products where customer clearly intended to buy but scheme was mentioned first - analyze if the purchase decision was truly influenced by the scheme or not]
    - [Note: These are products customer buys based on regular demand/habit/necessity]

B. Scheme Based Orders (Customer commits to buy ONLY BECAUSE schemes influenced their decision)
    - [List products where customer showed hesitation, said no initially, or was undecided BUT changed their mind specifically because of the scheme/offer]
    - [Include products where customer increased quantity due to schemes]
    - [Note: These purchases were clearly driven by the schemes/offers - customer behavior changed due to the incentive]

CRITICAL ANALYSIS REQUIRED - Look for these indicators:

**For Regular Buying:**
- Customer asks for the product immediately without hearing schemes
- Customer says "I need this" or "Give me [quantity]" before schemes are mentioned
- Customer shows clear intent to purchase regardless of offers
- Customer maintains same quantity even after hearing schemes

**For Scheme-Based Buying:**
- Customer initially hesitates or says "Let me think" but changes mind after scheme
- Customer says "No" first but then says "Okay" after hearing the offer
- Customer increases quantity specifically for the scheme (e.g., "Then give me 10kg instead of 5kg")
- Customer explicitly mentions the scheme as reason for buying (e.g., "Because of the free piece, I'll take it")
- Customer compares and decides based on the offer value

**IMPORTANT:** If customer was already planning to buy and scheme was just mentioned coincidentally, categorize as REGULAR buying, not scheme-based.
------------------------------------------------------------

4. Competitive Intelligence & Customer Psychology

A. Competitor Brand Analysis
For EACH competitor brand and product mentioned, document separately:
    for example a same brand may have multiple products mentioned in the conversation, DONT combine them and document each product separately and Multiple brands may be mentioned for same product also, DONT combine them and document each brand separately.

**Brand 1:**
- Brand Name: [e.g., Shakti, Nandi, etc.]
- Product: Which product categories?
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

**Brand 1:**
- Brand Name: [Same Brand as Previous]
- Product: Which product categories? (Another product of same brand)
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

**Brand 2:** (Continue for each additional competitor brand mentioned until all are covered)
- Brand Name: [Next competitor brand]
- Product: Which product categories?
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

Example:

**Brand 1:**
- Brand Name: Shakti
- Product: Gothumai Maavu (Wheat Flour)
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

**Brand 2:**
- Brand Name: Shakti
- Product: Rava Maavu
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

**Brand 3:**
- Brand Name: Nandi
- Product: Rava Maavu
- Customer's Current Status: Does customer stock it? How much?
- Reasons for Preference: Why does customer prefer this brand than Naga in detail?
    - [Price? Consumers Choice? Taste? Local brand? Habit? Promotions? etc..]
- Category:
    - [Based on the reason Categorize whether it is due to Price Concern or Local Brand Preference or Taste & Quality Preference or Brand Loyalty / Trust or Availability & Supply Strength or Retailer Margin Advantage or Promotions & Schemes or Packaging Preference or Customer Demand / Pull or Regional Taste / Cultural Fit or Other Factors]
      IMPORTANT - Choose the Category only on the list of reasons mentioned above, dont change the list.

B. Online Retailers Mentioned
For EACH online retailer mentioned, document separately:

**Retailer 1:**
- Name: [e.g., Amazon]
- Product Range: What products do they offer?
- Pricing Strategy: How do their prices compare to Naga?
- Customer Perception: How do customers view this retailer?
- Unique Selling Points: What makes this retailer stand out?

**Retailer 2:**
- Name: [Next online retailer]
- Product Range: What products do they offer?
- Pricing Strategy: How do their prices compare to Naga?
- Customer Perception: How do customers view this retailer?
- Unique Selling Points: What makes this retailer stand out?

C. Customer Buying Psychology
- What truly drives purchase decisions? (rank by importance)
- Is it price, brand recognition, customer demand, margins, or something else?
- Customer's risk tolerance (willing to try new brands?)
- Stock rotation preferences (fast-moving vs slow-moving)
- Is he open to switching if Naga offers better schemes or prices?
- How is the customer buying behaviour? (Like is he buys product if more schemes are offered, or if the product is on discount, or if the product is a well-known brand, or if more free pieces are offered, etc.)

------------------------------------------------------------

5. Salesperson Effectiveness Score:

Based on specific criteria - Score each component objectively.  

IMPORTANT: If any criterion does not apply to this conversation (e.g., no competitor brands mentioned → Competitor handling = N/A), then:
1. Mark that category as "N/A".
2. Give full score for that category.

---

**Product promotion (30% weight):** _/10
- 8-10: Presented 5+ Naga products with clear benefits and schemes
- 6-7: Presented 3-4 Naga products adequately  
- 4-5: Presented 1-2 Naga products with limited detail
- 1-3: Minimal product presentation

**Scheme leverage (20% weight):** _/10
- 8-10: Actively promoted multiple schemes and free offers
- 6-7: Mentioned some schemes but didn't emphasize strongly
- 4-5: Basic mention of schemes without detail
- 1-3: No schemes mentioned or poorly explained

**Competitor handling (25% weight):** _/10
- 8-10: Directly addressed competitor advantages with counter-arguments
- 6-7: Acknowledged competitors but weak counter-positioning
- 4-5: Mentioned competitors but didn't address customer concerns
- 1-3: Failed to address competitive threats

**Customer psychology understanding (25% weight):** _/10
- 8-10: Clearly understood customer's priorities and adapted pitch accordingly
- 6-7: Showed some understanding of customer needs
- 4-5: Basic awareness of customer concerns
- 1-3: Poor understanding of what drives customer decisions

---

**Final Score Calculation:**
- If all 4 criteria apply:  
  (Product promotion × 0.3) + (Scheme leverage × 0.2) + (Competitor handling × 0.25) + (Customer psychology × 0.25)  

------------------------------------------------------------

6. Salesperson Ability Analysis
- How salesperson handles the conversation, objections, competitor mentions, and customer concerns.

------------------------------------------------------------

7. Product Price Analysis
 What are all the Naga products that the customer thinks the price is too high?
 If so, list them with details like which product, what price point, and customer's exact concerns.

------------------------------------------------------------

8. Salesperson Strengths
- [Strength 1]
- [Strength 2]
- [Strength 3]

------------------------------------------------------------

9. Areas for Improvement
- [Improvement 1]
- [Improvement 2]
- [Improvement 3]

------------------------------------------------------------

ANALYSIS RULES

✓ Extract ALL numeric data (quantities, prices, pack sizes, margins, discounts)
✓ Clearly separate Naga products from ALL competitor brands
✓ Note EVERY competitor brand name mentioned - don't group them generically
✓ Capture the customer's REAL reasons for preferences (not just what they say on surface)
✓ Identify psychological factors beyond price (brand loyalty, habit, risk aversion, etc.)
✓ Highlight cases where customer prefers competitor DESPITE Naga advantages
✓ Document local/regional brand dynamics and home-ground advantages
✓ Assess whether salesperson understood the customer's true concerns
✓ Provide actionable, specific recommendations - not generic advice
✓ Use a dynamic matrix - only include categories relevant to THIS conversation
✓ Analyze audio by using both transcription as well as audio, as sometimes transcription might help in understanding the context better and audio might give better clarity on tone, emphasis, pauses, etc.
✓ Clearly Identify the Salesperson and the Customer in the conversation before analysis by the conversation context.
✓ Identify in what context does the salesperson is addressing about the price of the Naga product is high or low than the competitors before adding it to the price analysis section or at some other sections, Because sometimes the salesperson might be saying that the price is low compared to competitors, not high so it should not be added to the price analysis section as a high price concern. These things should be addressed carefully.
✓ IMPORTANT: The Analysis report should be in English only.
✓ Always cross-check the speaker before assigning brand ownership:
  - Sales Rep statements = Naga context
  - Customer statements = Competitor context
  - Never assume ownership without confirming who said it.
------------------------------------------------------------

Brand & Product Mapping
 
# Speaker & Brand Context Mapping
| Speaker | Brand | Product | Context / Ownership Confirmation |
|----------|--------|----------|--------------------------------|
| Sales Rep | Naga | [Product] | Confirmed as Naga (own product) |
| Customer | [Brand] | [Product] | Competitor product mentioned by customer |
| Sales Rep | — | [Product] | No brand mentioned → assumed Naga |
| Customer | — | [Product] | Ambiguous – Needs context |

CONSISTENCY REQUIREMENTS

For SCORING: Use the exact scoring rubric provided. Base scores on objective evidence from the conversation, not subjective impressions.

For ANALYSIS: Focus on factual observations. Use specific quotes and examples from the conversation rather than generalizations.

For RECOMMENDATIONS: Base suggestions on specific gaps identified in the conversation, not generic sales advice.

------------------------------------------------------------

------------------------------------------------------------

MANDATORY JSON OUTPUT FORMAT - FOLLOW THIS EXACT STRUCTURE

You MUST return your response as a valid JSON object following this exact structure.
This structure preserves ALL headings and subheadings from the analysis report:

```json
{
  "brand_product_mapping": {
    "naga_brand_products": {
      "products_list": []
    },
    "competitor_brands_mentioned": []
  },
  "conversation_summary": {
    "summary_points": []
  },
  "sales_matrix": {
    "naga_products_performance": {
      "naga_products_promoted": "",
      "volume_pushed_upselling": "",
      "schemes_offered": {
        "description": "",
        "scheme_details": []
      },
      "cross_selling_within_naga_portfolio": "",
      "acceptance_rejection": {
        "accepted": [],
        "rejected": []
      }
    },
    "sales_barriers": {
      "objections_raised": "",
      "competitor_advantages_cited": ""
    }
  },
  "customer_buying_patterns": {
    "regularly_buying_products": {
      "description": "Customer commits to buy BEFORE schemes are explained OR shows clear intent to buy regardless of schemes",
      "products": []
    },
    "scheme_based_orders": {
      "description": "Customer commits to buy ONLY BECAUSE schemes influenced their decision",
      "products": []
    }
  },
  "competitive_intelligence_and_customer_psychology": {
    "competitor_brand_analysis": [
      {
        "brand_name": "",
        "product": "",
        "customer_current_status": "",
        "reasons_for_preference": "",
        "category": ""
      }
    ],
    "online_retailers_mentioned": [
      {
        "name": "",
        "product_range": "",
        "pricing_strategy": "",
        "customer_perception": "",
        "unique_selling_points": ""
      }
    ],
    "customer_buying_psychology": {
      "purchase_decision_drivers_ranked": [],
      "risk_tolerance": "",
      "stock_rotation_preferences": "",
      "openness_to_switching": "",
      "buying_behaviour": ""
    }
  },
  "salesperson_effectiveness_score": {
    "scores": {
      "product_promotion": {
        "score": 0,
        "weight_percentage": 30,
        "justification": ""
      },
      "scheme_leverage": {
        "score": 0,
        "weight_percentage": 20,
        "justification": ""
      },
      "competitor_handling": {
        "score": 0,
        "weight_percentage": 25,
        "justification": "",
        "is_na": false
      },
      "customer_psychology_understanding": {
        "score": 0,
        "weight_percentage": 25,
        "justification": ""
      }
    },
    "final_score_calculation": {
      "formula": "",
      "final_score": 0
    }
  },
  "salesperson_ability_analysis": "",
  "product_price_analysis": {
    "summary": "",
    "high_price_products": [
      {
        "product": "",
        "price_point": "",
        "customer_exact_concerns": ""
      }
    ]
  },
  "salesperson_strengths": [],
  "areas_for_improvement": []
}
```

CRITICAL JSON FORMATTING RULES:

1. **Return ONLY the JSON object** - no markdown code blocks (no ```json), no explanatory text before or after
2. **Ensure all strings are properly escaped** - use \" for quotes inside strings
3. **Use arrays ([])** for all lists of items
4. **Use empty strings ("")** for text fields with no data
5. **Use empty arrays ([])** for list fields with no data
6. **All numeric scores must be actual numbers**, not strings (e.g., 9 not "9")
7. **Maintain proper JSON syntax** - no trailing commas, proper quotation marks
8. **For the "category" field** in competitor analysis, use ONLY these exact values:
   - "Price Concern"
   - "Discount Concern"
   - "Product Variety"
   - "Product Package Size"
   - "Other factors"
9. **For boolean fields** like "is_na", use true/false (not "true"/"false")
10. **Product lists should be detailed strings or objects** with full context as shown in the example document

FIELD MAPPING GUIDE (Ensure all content fits properly):

**brand_product_mapping.naga_brand_products.products_list**: 
- Include ALL Naga products mentioned with sizes/variants (e.g., "Rava (200g, 500g, 1kg)")

**competitor_brands_mentioned**: 
- List each competitor with context (e.g., "MTR: Mentioned briefly by the customer to someone else in the store, not as a direct competitor in the conversation")

**sales_matrix.naga_products_performance.schemes_offered.scheme_details**:
- Array of objects, each with product name and scheme details
- Example: [{"product": "Rava (500g & 1kg bags)", "scheme": "Get 6 packets of 500g Atta free"}]

**sales_matrix.naga_products_performance.acceptance_rejection**:
- "accepted": Array of products with quantities (e.g., "Rava (1 bag of 1kg, 1 bag of 500g)")
- "rejected": Array of rejected products with reasons if available

**customer_buying_patterns**: 
- Each product entry should include the reasoning/evidence for categorization

**salesperson_effectiveness_score.scores.competitor_handling.is_na**:
- Set to true if no competitors were discussed, false otherwise
- If true, explain in justification field

CRITICAL REMINDERS

- DO NOT assume any brand is Naga unless explicitly stated
- DO NOT group competitors as "other brands" – name each specifically in the JSON array
- DO capture both stated reasons AND underlying psychology
- DO identify non-price factors driving brand preference
- DO note when customer prefers competitor despite Naga being cheaper/better
- **RETURN ONLY VALID JSON - NO MARKDOWN CODE BLOCKS, NO ADDITIONAL TEXT OR FORMATTING**
- **START your response directly with the opening brace { and end with the closing brace }**
"""
    
    response = model.generate_content([
        analysis_prompt,
        {"mime_type": "audio/mp3", "data": audio_file}
    ])
    
    return response.text

# Streamlit app
def main():
    st.set_page_config(
        page_title="Sales Call Analyzer",
        page_icon="📊",
        layout="wide"
    )

    # Initialize page state
    if 'page' not in st.session_state:
        st.session_state['page'] = 'home'

    def render_dashboard():
        st.title("Sales Performance Dashboard")

        excel_path = os.path.join("data", 'monthly.xlsx')

        # Load Data
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return

        if 'Period' not in df.columns:
            st.error("❌ 'Period' column not found in the data.")
            return
        
        Months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
        selected_month = st.selectbox("Select Month", Months)

        # Filter DataFrame based on selected month
        df = df[df['Period'] == selected_month]

        if df.empty:
            st.warning(f"No data found for the selected month: {selected_month}")
            return

        # --- KPIs ---
        st.subheader(f"{selected_month} Sales Performance Overview")
        kpi_cols = st.columns(4)
        kpi_cols[0].metric("🧾 Total Reports", f"{df['Total Reports Analysed'].iloc[0]}")
        kpi_cols[1].metric("🛒 Overall Sales Effectiveness", f"{df['Overall Sales Effectiveness'].iloc[0]}")
        kpi_cols[2].metric("☎️ Total Duration", f"{df['Total Duration'].iloc[0]}")
        kpi_cols[3].metric("📞 Average Call Duration", f"{df['Average Duration'].iloc[0]}")

        st.divider()

        # ========================
        # PRODUCT DISCUSSION TREEMAP
        # ========================
        if 'Products Discussed' in df.columns:
            st.subheader("Naga Product Mention Rate")

            try:
                counts = parse_explicit_counts(df["Products Discussed"])

                if counts:
                    freq_df = pd.DataFrame(counts.items(), columns=['Product', 'Count'])
                    freq_df = freq_df.sort_values(by='Count', ascending=False)

                    fig = px.treemap(
                        freq_df,
                        path=[px.Constant("Product Mention Rate"), 'Product'],
                        values='Count',
                        color='Count',
                        color_continuous_scale='Blues',
                    )
                    fig.update_layout(
                        margin=dict(t=50, l=25, r=25, b=25),
                        uniformtext=dict(minsize=10, mode='hide')
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No product discussion data found.")
            except Exception as e:
                st.warning(f"Could not generate product treemap: {e}")
        else:
            st.info("The column 'Products Discussed' was not found in the Excel file.")

        st.divider()

        # ========================
        # COMPETITOR TREEMAP
        # ========================
        if 'Competitors' in df.columns:
            st.subheader("Competitor Mention Rate")

            try:
                counts = parse_explicit_counts(df["Competitors"])
                
                if counts:
                    freq_df = pd.DataFrame(counts.items(), columns=['Competitor', 'Count'])
                    freq_df = freq_df.sort_values(by='Count', ascending=False)

                    fig = px.treemap(
                        freq_df,
                        path=[px.Constant("Competitor Mention Rate"), 'Competitor'],
                        values='Count',
                        color='Count',
                        color_continuous_scale='Blues',
                    )
                    fig.update_layout(
                        margin=dict(t=50, l=25, r=25, b=25),
                        uniformtext=dict(minsize=10, mode='hide')
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No competitor data found.")
            except Exception as e:
                st.warning(f"Could not generate competitor treemap: {e}")
        else:
            st.info("The column 'Competitors' was not found in the Excel file.")

        st.divider()

        # ========================
        # COMPETITOR PRODUCT TREEMAP
        # ========================
        if 'Competitor Products' in df.columns:
            st.subheader("Competitor Product Preference")

            try:
                counts = parse_explicit_counts(df["Competitor Products"])
                
                if counts:
                    freq_df = pd.DataFrame(counts.items(), columns=['Product', 'Count'])
                    freq_df = freq_df.sort_values(by='Count', ascending=False)

                    fig = px.treemap(
                        freq_df,
                        path=[px.Constant("Competitor Product Preference"), 'Product'],
                        values='Count',
                        color='Count',
                        color_continuous_scale='Blues',
                    )
                    fig.update_layout(
                        margin=dict(t=50, l=25, r=25, b=25),
                        uniformtext=dict(minsize=10, mode='hide')
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No competitor product data found.")
            except Exception as e:
                st.warning(f"Could not generate competitor product treemap: {e}")
        else:
            st.info("The column 'Competitor Products' was not found in the Excel file.")
        
        st.divider()

        # ========================
        # PRICING CONCERN TREEMAP
        # ========================
        # if 'Pricing Concerns' in df.columns:
        #     st.subheader("Products with Pricing Concerns")

        #     try:
        #         all_concerns = []
        #         for _, row in df.iterrows():
        #             concerns = [p.strip() for p in str(row["Pricing Concerns"]).split(",") if p.strip()]
        #             all_concerns.extend(concerns)
        #         if all_concerns:
        #             counts = Counter(all_concerns)
        #             freq_df = pd.DataFrame(counts.items(), columns=['Concern', 'Count'])
        #             freq_df = freq_df.sort_values(by='Count', ascending=False)

        #             fig = px.treemap(
        #                 freq_df,
        #                 path=[px.Constant("Products with Pricing Concerns"), 'Concern'],
        #                 values='Count',
        #                 color='Count',
        #                 color_continuous_scale='Blues',
        #             )
        #             fig.update_layout(
        #                 margin=dict(t=50, l=25, r=25, b=25),
        #                 uniformtext=dict(minsize=10, mode='hide')
        #             )
        #             st.plotly_chart(fig, use_container_width=True)
        #         else:
        #             st.info("No pricing concern data found.")
        #     except Exception as e:
        #         st.warning(f"Could not generate pricing concern treemap: {e}")
        # else:
        #     st.info("The column 'Pricing Concerns' was not found in the Excel file.")

    def render_individual_dashboard():
        st.title("Individual Salesperson Dashboard")

        excel_path = os.path.join("data", 'individually.xlsx')

        # Load Data
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return

        # Ensure salesperson column exists
        if 'SalesPerson' not in df.columns:
            st.error("❌ 'SalesPerson' column not found in the data.")
            return

        # Dropdown to select salesperson
        salesperson_names = sorted(df['SalesPerson'].dropna().unique())
        selected_salesperson = st.selectbox("Select Salesperson", salesperson_names)

        # Filter data for the selected salesperson
        person_df = df[df['SalesPerson'] == selected_salesperson]

        if person_df.empty:
            st.warning(f"No data found for salesperson: {selected_salesperson}")
            return

        # --- KPIs ---
        st.subheader(f"Performance Overview — {selected_salesperson}")
        kpi_cols = st.columns(4)
        kpi_cols[0].metric("🧾 Total Reports", f"{person_df['Total Reports Analysed'].iloc[0]}")
        kpi_cols[1].metric("🛒 Sales Effectiveness", f"{person_df['Overall Sales Effectiveness'].iloc[0]}")
        kpi_cols[2].metric("☎️ Total Duration", f"{person_df['Total Duration'].iloc[0]}")
        kpi_cols[3].metric("📞 Average Call Duration", f"{person_df['Average Duration'].iloc[0]}")

        st.divider()

        score_columns = ['Product promotion', 'Scheme leverage', 'Competitor handling', 'Customer psychology understanding']
        avg_scores = person_df[score_columns].iloc[0].tolist()
        categories = ['Product Promotion Skill', 'Scheme Utilization', 'Competitor Handling Skill', 'Customer Understanding']

        # Create Radar chart
        fig = go.Figure(data=go.Scatterpolar(
            r=avg_scores + [avg_scores[0]], 
            theta=categories + [categories[0]],
            fill='toself',
            name='Average Monthly Scores',
            line_color="#6873f9",
            fillcolor='rgba(164, 173, 248)'
        ))

        # Layout settings
        fig.update_layout(
            polar=dict(
                radialaxis=dict(visible=True, range=[0,10]), bgcolor='#e5ecf6',
                angularaxis=dict(tickfont=dict(size=16))
                ),
            showlegend=False,
            height = 600
        )
        st.subheader(f"Performance Breakdown")
        st.plotly_chart(fig, use_container_width=True)

        st.divider()

        excel_path = os.path.join("data", 'TopSalesPitch.xlsx')
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return
        st.subheader("Most Pitched Items")
        st.bar_chart(df, x="Product Name", y="Mention Count", x_label="Product Name", y_label="Mention Count", height=500, sort="-Mention Count")

        st.divider()

        excel_path = os.path.join("data", 'LeastPitchedItems.xlsx')
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return
        st.subheader("Least Pitched Items")
        st.bar_chart(df, x="Product Name", y="Pitch Count", x_label="Product Name", y_label="Pitch Count", height=500, sort="Pitch Count")
    
    def summary_dashboard():
        
        st.title("Summary Dashboard")
        monthly_excel_path = os.path.join("data", 'monthly.xlsx')

        st.divider()

        # Load Data
        try:
            m_df = pd.read_excel(monthly_excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return

        individual_excel_path = os.path.join("data", 'individually.xlsx')

        # Load Data
        try:
            p_df = pd.read_excel(individual_excel_path)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            if st.button("⬅️ Back to Home"):
                st.session_state['page'] = 'home'
            return
        
        score_column = 'Overall Sales Effectiveness'

        # Categorize scores
        def categorize_score(score):
            if score > 8:
                return 'Well'
            elif 6 <= score <= 8:
                return 'Moderate'
            else:
                return 'Poor'

        p_df['Performance Category'] = p_df[score_column].apply(categorize_score)

        # Count the number of people in each category
        category_counts = p_df['Performance Category'].value_counts().reset_index()
        category_counts.columns = ['Category', 'Count']

        # Create pie chart
        fig = px.pie(
            category_counts,
            names='Category',
            values='Count',
            color='Category',
            color_discrete_map={
                'Well': '#2ca02c',        # green
                'Moderate': '#ffbf00',    # yellow
                'Poor': '#d62728'         # red
            },
            height = 600
        )

        # 1. Increase the font size of the labels displayed *on* the pie slices
        fig.update_traces(
            textfont_size=18, # Adjust this value (e.g., 14, 18, 20) as needed
            marker=dict(line=dict(color='#000000', width=1)) # Optional: adds a slight border for visibility
        )

        # 2. Optional: Increase the font size of the *legend* entries
        fig.update_layout(
            legend=dict(
                font=dict(size=14) # Adjust this value as needed
            )
        )

        st.subheader("Overall Sales Performance Distribution")
        # Display the chart
        st.plotly_chart(fig, use_container_width=True)
        
        st.divider()

        Months = ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
        selected_month = st.selectbox("Select Month", Months)

        # Filter DataFrame based on selected month
        m_df = m_df[m_df['Period'] == selected_month]
        p_count = c_count = cp_count = pc_count = 0

        if 'Products Discussed' in m_df.columns:
            try:
                counts = parse_explicit_counts(m_df["Products Discussed"])
                p_count = sum(counts.values())
            except Exception as e:
                st.warning("Couldn't process Product Discussed column")
        
        if 'Competitors' in m_df.columns:
            try:
                counts = parse_explicit_counts(m_df["Competitors"])
                c_count = sum(counts.values())
            except Exception as e:
                st.warning("Couldn't process Competitors column")
        
        if 'Competitor Products' in m_df.columns:
            try:
                counts = parse_explicit_counts(m_df["Competitor Products"])
                cp_count = sum(counts.values())
            except Exception as e:
                st.warning("Couldn't process Competitor Products column")
        
        if 'Pricing Concerns' in m_df.columns:
            try:
                counts = parse_explicit_counts(m_df["Pricing Concerns"])
                pc_count = sum(counts.values())
            except Exception as e:
                st.warning("Couldn't process Pricing Concerns column")
        # --- Summary of discussion counts ---
        st.subheader("Overall Discussion Summary")

        # Collect the counts into a dataframe
        summary_data = pd.DataFrame({
            'Category': ['Products Discussed', 'Competitors', 'Competitor Products', 'Pricing Concerns'],
            'Count': [p_count, c_count, cp_count, pc_count]
        })

        # Create a bar chart
        fig_summary = px.bar(
            summary_data,
            x='Category',
            y='Count',
            text='Count',
            color='Count',
            color_continuous_scale='Bluered',
            height = 500
        )

        # Style adjustments
        fig_summary.update_traces(textposition='outside', textfont_size=15)
        fig_summary.update_layout(
            xaxis_title="Discussion Category",
            yaxis_title="Total Mentions",
            template='simple_white',
            yaxis=dict(showgrid=True, zeroline=False, title_font=dict(size=16), tickfont=dict(size=14)), 
            font = dict(size=14),
            xaxis = dict(title_font=dict(size=16), tickfont=dict(size=14)),
        )

        # Display in Streamlit
        st.plotly_chart(fig_summary, use_container_width=True)

        st.divider()

        # # Path to your Excel file
        # excel_path = os.path.join("data", "products.xlsx")

        # try:
        #     df = pd.read_excel(excel_path)
            
        #     # Optional: Clean column names
        #     df.columns = [col.strip().title() for col in df.columns]

        #     st.subheader("Product-wise Competitor Overview")

        #     # Display the dataframe neatly
        #     st.dataframe(
        #         df.style.set_properties(**{
        #             'background-color': '#f9f9f9',
        #             'color': '#333',
        #             'border-color': '#ddd'
        #         }),
        #         use_container_width=True,
        #         hide_index=True
        #     )

        # except Exception as e:
        #     st.error(f"❌ Failed to load product data: {e}")

    def competitor_performance():
        st.title("Competitor Performance Analysis")

        excel_path = os.path.join("data", "products.xlsx")

        # Load Excel
        try:
            df = pd.read_excel(excel_path)
        except Exception as e:
            st.error(f"❌ Failed to load data file: {e}")
            return

        # Validate necessary columns
        required_columns = ['Products', 'Potential Competitors', 'Reason']
        if not all(col in df.columns for col in required_columns):
            st.error(f"❌ Missing required columns. Expected: {required_columns}")
            return

        # Clean and prepare data
        df = df.dropna(subset=required_columns)
        for col in required_columns:
            if df[col].dtype == 'object':
                df[col] = df[col].astype(str).str.strip()

        # Dropdown to select product
        selected_product = st.selectbox(
            "🛒 Select a Product",
            sorted(df['Products'].unique())
        )

        # Filter data for selected product
        filtered_df = df[df['Products'] == selected_product]

        if filtered_df.empty:
            st.warning("No data available for the selected product.")
            return

        # Parse comma-separated competitors and reasons
        expanded_rows = []
        
        for _, row in filtered_df.iterrows():
            # Split competitors and reasons by comma
            competitors = [c.strip() for c in str(row['Potential Competitors']).split(',')]
            reasons = [r.strip() for r in str(row['Reason']).split(',')]
            
            # Create a row for each competitor-reason pair
            # Match them by index (assuming they correspond positionally)
            for i, competitor in enumerate(competitors):
                if i < len(reasons):  # Ensure we have a corresponding reason
                    expanded_rows.append({
                        'Potential Competitors': competitor,
                        'Reason': reasons[i]
                    })

        # Convert to DataFrame
        expanded_df = pd.DataFrame(expanded_rows)
        
        if expanded_df.empty:
            st.warning(f"No competitor data available for {selected_product}.")
            return

        # Count occurrences of each competitor-reason combination
        count_df = (
            expanded_df.groupby(['Potential Competitors', 'Reason'])
            .size()
            .reset_index(name='Count')
        )

        # Sort by competitor and count for better visualization
        count_df = count_df.sort_values(['Potential Competitors', 'Count'], ascending=[True, False])

        # Create stacked bar chart
        fig = px.bar(
            count_df,
            x='Potential Competitors',
            y='Count',
            color='Reason',
            title=f"Competitor Performance for {selected_product}",
            text='Count',
            barmode='stack',
            color_discrete_sequence=px.colors.qualitative.Set3, 
            height=500
        )

        fig.update_traces(textposition='inside', textfont_size=15)
        
        fig.update_layout(
            xaxis_title="Potential Competitors",
            yaxis_title="Count of Mentions",
            legend_title="Reasons",
            plot_bgcolor="#f9f9f9",
            paper_bgcolor="#ffffff",
            font=dict(size=13),
            title_x=0.5,
            xaxis = dict(title_font=dict(size=16), tickfont=dict(size=14), categoryorder='total descending'),
            yaxis = dict(title_font=dict(size=16), tickfont=dict(size=14)),
        )

        st.plotly_chart(fig, use_container_width=True)  
    
    def product_performance():
        st.title("Product Pain-Point Analytics")

        # Load Excel
        file_path = os.path.join("data", "concerns.xlsx")
        try:
            df = pd.read_excel(file_path)
        except Exception as e:
            st.error(f"Failed to load file: {e}")
            return
        
        # Validate columns
        if "Products" not in df.columns or "Concerns" not in df.columns:
            st.error("Excel must contain 'Product' and 'Concerns' columns")
            return

        # Product dropdown
        products = sorted(df["Products"].dropna().unique())
        selected_product = st.selectbox("Select Product", products)

        # Get concerns list for selected product
        concerns_series = df[df["Products"] == selected_product]["Concerns"]
        concern_counts = parse_explicit_counts(concerns_series)

        # Convert to DataFrame
        concern_df = pd.DataFrame(list(concern_counts.items()), columns=["Concern", "Count"])
        concern_df = concern_df.sort_values(by="Count", ascending=False)

        st.subheader(f"Key Concern Areas for {selected_product}")

        # Bar chart
        fig = px.bar(
            concern_df,
            x="Concern",
            y="Count",
            text="Count",
            color="Concern",
            color_discrete_sequence=px.colors.qualitative.Set2, 
            height=600
        )
        fig.update_traces(textposition="outside", textfont_size=15)
        fig.update_layout(
            xaxis_title="Concern Type",
            yaxis_title="Frequency",
            showlegend=True,
            template="simple_white",
            xaxis = dict(title_font=dict(size=16), tickfont=dict(size=14)),
            yaxis = dict(title_font=dict(size=16), tickfont=dict(size=14))
        )

        st.plotly_chart(fig, use_container_width=True)

    # Sidebar for instructions and navigation
    with st.sidebar:
        
        if st.button("🏡 Home", width="stretch"):
            st.session_state['page'] = 'home'
            st.rerun()

        # if st.button("📋 Summary Dashboard", width="stretch"):
        #     st.session_state['page'] = 'summary_dashboard'
        #     st.rerun()

        if st.button("💰 Sales Performance", width="stretch"):
            st.session_state['page'] = 'dashboard'
            st.rerun()

        if st.button("👥 Salesperson Dashboard", width="stretch"):
            st.session_state['page'] = 'individual_dashboard'
            st.rerun()

        # if st.button("🎯 Competitor Performance Dashboard", width="stretch"):
        #     st.session_state['page'] = 'competitor_performance'
        #     st.rerun()

        if st.button("⚠️ Product Pain-Point Analytics", width="stretch"):
            st.session_state['page'] = 'product_performance'
            st.rerun()

    # Route pages
    if st.session_state.get('page', 'home') == 'dashboard':
        render_dashboard()
        return
    
    if st.session_state.get('page', 'home') == 'individual_dashboard':
        render_individual_dashboard()
        return

    # if st.session_state.get('page', 'home') == 'summary_dashboard':
    #     summary_dashboard()
    #     return
    
    # if st.session_state.get('page', 'home') == 'competitor_performance':
    #     competitor_performance()
    #     return
    
    if st.session_state.get('page', 'home') == 'product_performance':
        product_performance()
        return
    
    st.title("Sales Call Analyzer")
    st.divider()
    salespersonName = st.text_input("Salesperson Name")
    storeName = st.text_input("Store Name")
    st.divider()
    # Main content area (home)
    col1, col2 = st.columns([1, 2])

    with col1:
        st.header("📁 Upload Audio")

        # Audio file uploader
        uploaded_file = st.file_uploader(
            "Choose an audio file",
            type=['mp3', 'wav', 'mp4', 'm4a', 'ogg'],
            help="Upload your sales conversation audio file"
        )

        if uploaded_file is not None:
            st.success(f"✅ File uploaded: {uploaded_file.name}")

            # Audio player
            st.audio(uploaded_file)

            # Analyze button
            if st.button("Analyze Audio", type="primary"):
                with st.spinner("🔄 Analyzing audio..."):
                    try:
                        # Read the uploaded file
                        audio_data = uploaded_file.read()

                        # Analyze with Gemini
                        analysis = analyze_audio_with_gemini(audio_data)
                        print("Analysis Result ::::::", analysis)

                        # Clean the response if it contains markdown code blocks
                        clean_analysis = analysis.strip()
                        if clean_analysis.startswith("```json"):
                            clean_analysis = clean_analysis[7:]
                        elif clean_analysis.startswith("```"):
                            clean_analysis = clean_analysis[3:]
                        if clean_analysis.endswith("```"):
                            clean_analysis = clean_analysis[:-3]
                        clean_analysis = clean_analysis.strip()

                        if not clean_analysis:
                             st.error("The model returned an empty response.")
                        else:
                             report = convert_sales_report_to_string(json.loads(clean_analysis))
                             
                             # Store in session state
                             st.session_state['analysis_result'] = report

                        st.success("✅ Analysis completed!")

                    except Exception as e:
                        st.error(f"❌ Error analyzing audio: {str(e)}")

    with col2:
        st.header("Analysis Results")

        if 'analysis_result' in st.session_state:

            if st.button("Clear Analysis"):
                if 'analysis_result' in st.session_state:
                    del st.session_state['analysis_result']
                st.rerun()

            # Display analysis in a nice format
            st.markdown("### Sales Performance Analysis")

            # Create tabs for better organization
            tab1, tab2 = st.tabs(["📋 Full Report", "💾 Export"])

            with tab1:
                # Display the analysis with proper formatting
                analysis_text = st.session_state['analysis_result']
                st.markdown(analysis_text)

            with tab2:
                # Download button for the analysis as Word document
                analysis_text = st.session_state['analysis_result']

                # Create Word document
                doc = Document()
                doc.add_heading('Sales Performance Analysis Report', 0)

                # Process content to handle markdown formatting
                def process_line_to_word(doc, line):
                    line = line.strip()
                    if not line:
                        return

                    # Check if it's a heading (starts with #)
                    if line.startswith('#'):
                        # Count the number of # to determine heading level
                        heading_level = 0
                        for char in line:
                            if char == '#':
                                heading_level += 1
                            else:
                                break

                        # Remove the # symbols and add as heading
                        heading_text = line.lstrip('#').strip()
                        if heading_text:
                            doc.add_heading(heading_text, min(heading_level, 9))
                    else:
                        # Handle regular paragraphs with bold formatting
                        paragraph = doc.add_paragraph()

                        # Split text by ** to handle bold formatting
                        parts = line.split('**')

                        for i, part in enumerate(parts):
                            if part:  # Only add non-empty parts
                                if i % 2 == 0:  # Even index = normal text
                                    paragraph.add_run(part)
                                else:  # Odd index = bold text
                                    paragraph.add_run(part).bold = True

                # Add content to document
                for line in analysis_text.split('\n'):
                    process_line_to_word(doc, line)

                # Save to BytesIO
                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                # Remove file extension from uploaded file name for the report
                if uploaded_file is not None:
                    base_filename = os.path.splitext(uploaded_file.name)[0]
                else:
                    base_filename = "analysis"
                st.download_button(
                    label="📄 Download Analysis Report (Word)",
                    data=doc_buffer.getvalue(),
                    file_name=f"{base_filename}_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        else:
            st.info("👆 Upload an audio file and click 'Analyze Audio' to see results here.")

if __name__ == "__main__":
    main()
